from selenium import webdriver
import time
import datetime as dt
import csv
import os
from selenium.common.exceptions import StaleElementReferenceException
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.common.by import By
from selenium.common.exceptions import TimeoutException
from selenium.webdriver.common.action_chains import ActionChains
from selenium.webdriver.common.keys import Keys
from bs4 import BeautifulSoup as soup

import docx
from docx.shared import Inches
from lxml import etree as ET
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def add_hyperlink(paragraph, url, text):
    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :return: The hyperlink object
    """

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink




def find_main_icon(username):
    os.chdir('Photos')
    os.chdir('Icons')
    os.chdir('MainTweeterIcons')
    photo_search = []
    for f in os.listdir():
        if username[1:]+'-F' in f:
            photo_search.append(f)
    os.chdir('..')
    os.chdir('..')
    os.chdir('..')
    for f in photo_search:
        if f[-4:]=='.png':
             return f
            
            
def find_quote_icon(username):
    os.chdir('Photos')
    os.chdir('Icons')
    os.chdir('QuoteIcons')
    photo_search = []
    for f in os.listdir():
        if username[1:]+'-F' in f:
            photo_search.append(f)
    os.chdir('..')
    os.chdir('..')
    os.chdir('..')
    for f in photo_search:
        if f[-4:]=='.png':
             return f
            
            
def find_tweet_photos(inum):
    if len(sheet[15][inum])>0:
        os.chdir('Photos')
        photo_name_list = []
        for f in os.listdir():
            if 'Tweet_'+str(inum)+'_Photo' in f and f[-4:]=='.png':
                photo_name_list.append(f)
        os.chdir('..')
        return photo_name_list
    
    
def find_quote_photos(inum):
    if len(sheet[33][inum])>0:
        os.chdir('Photos')
        os.chdir('QuotePhotos')
        photo_name_list = []
        for f in os.listdir():
            if 'Quote_'+str(inum)+'_Photo' in f and f[-4:]=='.png':
                photo_name_list.append(f)
        os.chdir('..')
        os.chdir('..')
        return photo_name_list
    
    
def find_article_pic(inum):
    os.chdir('Photos')
    os.chdir('ArticlePreviews')
    art_list = []
    for f in os.listdir():
        if 'Tweet_'+str(inum)+'_Article' in f and f[-4:]=='.png':
            art_list.append(f)
    os.chdir('..')
    os.chdir('..')
    try:
        return art_list[0]
    except:
        pass
            


o = open('Twitter_FratboyMarxist-2019-12-03_1349.csv','r',encoding='UTF-8')
r = csv.reader(o)
rows = []
for row in r:
    rows.append(row)
    
sheet = [list(t) for t in zip(*rows,)]


A = time.time()

doc = docx.Document()
doc.add_heading('Twitter Archive 2019-12-03',0)

for J in range(1,len(rows)):
    doc.add_heading('Tweet No. '+sheet[1][J],1)
    locals()['p1_%i'%J] = doc.add_paragraph()
    add_hyperlink(locals()['p1_%i'%J],sheet[8][J],'STATUS:  ')
    locals()['p1_%i'%J].add_run(sheet[0][J]+ '  [' + sheet[10][J] + '] ')

    if len(sheet[5][J])>0:
        locals()['p2_%i'%J] = doc.add_paragraph()
        locals()['p2_%i'%J].add_run(' - '+sheet[5][J]+' - '+sheet[6][J]+' - '+sheet[9][J])

    locals()['p3_%i'%J] = doc.add_paragraph()
    add_hyperlink(locals()['p3_%i'%J],'https://twitter.com/'+sheet[3][J][1:],'USER:  ')
    locals()['r3_%i'%J] = locals()['p3_%i'%J].add_run()
    locals()['r3_%i'%J].add_picture('Photos/Icons/MainTweeterIcons/'+find_main_icon(sheet[3][J]),height=Inches(0.35),width=Inches(0.35) )
    locals()['r3_%i'%J].add_text('  '+sheet[2][J]+' | '+sheet[3][J]+' | '+sheet[4][J])

    if len(sheet[11][J])>0:
        locals()['p4_%i'%J] = doc.add_paragraph()
        locals()['p4_%i'%J].add_run(sheet[11][J]).bold = True
        if len(sheet[12][J])>0:
            for i,L in enumerate(sheet[12][J].replace('[','').replace(']','').replace("'",'').replace(' ','').split(',')):
                add_hyperlink(locals()['p4_%i'%J],L,' [LNK%i] '%(i+1))

    if len(sheet[15][J])>0:
        locals()['p5_%i'%J] = doc.add_paragraph()
        locals()['p5_%i'%J].alignment = WD_ALIGN_PARAGRAPH.CENTER
        locals()['pics_layout_%i'%J] = locals()['p5_%i'%J].add_run()
        for p in find_tweet_photos(J):
            locals()['pics_layout_%i'%J].add_picture('Photos/'+p,width=Inches(1.5),height=Inches(1.5))

    if len(sheet[16][J])>0:
        locals()['p6_%i'%J] = doc.add_paragraph()
        add_hyperlink(locals()['p6_%i'%J],sheet[16][J],'VIDEO:  ')
        locals()['p6_%i'%J].add_run(sheet[18][J]+ ' long,  '+sheet[17][J])
        locals()['p6_%i'%J].alignment = WD_ALIGN_PARAGRAPH.CENTER


    if len(sheet[23][J])>0:
        locals()['p7_%i'%J] = doc.add_paragraph()
        locals()['p7_%i'%J].alignment = WD_ALIGN_PARAGRAPH.CENTER
        locals()['r7_%i'%J] = locals()['p7_%i'%J].add_run()
        try:
            locals()['r7_%i'%J].add_picture('Photos/ArticlePreviews/'+find_article_pic(J),width=Inches(4.2),height=Inches(2.1))
        except:
            pass
        locals()['art_tab_%i'%J] = doc.add_table(rows=3,cols=1,style='Light Grid')
        locals()['art_tab_%i'%J].cell(0,0).text = sheet[20][J]
        locals()['art_tab_%i'%J].cell(0,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        locals()['art_tab_%i'%J].cell(1,0).text = sheet[21][J]
        locals()['art_tab_%i'%J].cell(1,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        H = locals()['art_tab_%i'%J].cell(2,0).paragraphs[0] 
        locals()['art_tab_%i'%J].cell(2,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_hyperlink(H,sheet[23][J],sheet[22][J])

    if 'Quote ' in sheet[24][J]:
        locals()['p8_%i'%J] = doc.add_paragraph(style='Intense Quote')
        add_hyperlink(locals()['p8_%i'%J],'https://twitter.com/'+sheet[26][J][1:],'USER:  ')
        locals()['r8_%i'%J] = locals()['p8_%i'%J].add_run()
        locals()['r8_%i'%J].add_picture('Photos/Icons/QuoteIcons/'+find_quote_icon(sheet[26][J][1:]),height=Inches(0.35),width=Inches(0.35))
        locals()['r8_%i'%J].add_text('  '+sheet[25][J]+' | '+sheet[26][J]+' | '+sheet[27][J]).italic = False
        locals()['p9_%i'%J] = doc.add_paragraph(style='Intense Quote')
        locals()['r9_%i'%J] = locals()['p9_%i'%J].add_run()
        locals()['r9_%i'%J].add_text(' - '+sheet[29][J]+' - '+sheet[30][J]+' - '+'['+ sheet[31][J] +']').italic = False
        locals()['p10_%i'%J] = doc.add_paragraph(style='Intense Quote')
        locals()['r10_%i'%J] = locals()['p10_%i'%J].add_run()
        locals()['r10_%i'%J].add_text(sheet[32][J]).bold = False
        locals()['p11_%i'%J] = doc.add_paragraph()
        locals()['r11_%i'%J] = locals()['p11_%i'%J].add_run()
        try:
            for q in find_quote_photos(J):
                locals()['r11_%i'%J].add_picture('Photos/QuotePhotos/'+q,width=Inches(1.5),height=Inches(1.5))
        except:
            pass
        locals()['p11_%i'%J].alignment = WD_ALIGN_PARAGRAPH.CENTER
        locals()['p12_%i'%J] = doc.add_paragraph()
        if len(sheet[34][J])>0:
            add_hyperlink(locals()['p12_%i'%J],sheet[34][J],' QUOTE VIDEO POSTED ')
            locals()['p12_%i'%J].alignment = WD_ALIGN_PARAGRAPH.CENTER

    locals()['interactions_%i'%J] = doc.add_table(rows=1,cols=4,style='Light Grid')
    locals()['interactions_%i'%J].cell(0,0).text = sheet[35][J]
    locals()['interactions_%i'%J].cell(0,0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    locals()['interactions_%i'%J].cell(0,1).text = sheet[36][J]
    locals()['interactions_%i'%J].cell(0,1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    locals()['interactions_%i'%J].cell(0,2).text = sheet[37][J]
    locals()['interactions_%i'%J].cell(0,2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    locals()['interactions_%i'%J].cell(0,3).text = sheet[38][J]
    locals()['interactions_%i'%J].cell(0,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_page_break()
    
doc.save('ArchiveDocument.docx')
B = time.time()
print('Done!')
print(B-A, 'seconds')

